import os
import json
import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

KEYWORDS = [
    "gerar relatório", "exportar para word", "criar pdf", "salvar relatório", "gerar docx",
    "previsão de vendas", "forecast de vendas", "projetar receita", "prever faturamento",
    "análise de pipeline", "saúde do funil", "previsão de receita"
]

def create_professional_report(data, filename="Relatorio_Laura_AI.docx"):
    """Transforma dados de análise em um documento Word elegante."""
    doc = Document()
    
    # Estilo do Título Principal
    title = doc.add_heading('Auditoria Estratégica de Marketing & CRO', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Adicionar Meta Dados
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = meta.add_run(f"Data: {datetime.datetime.now().strftime('%d/%m/%Y %H:%M')}\n")
    run.font.size = Pt(10)
    if "url" in data:
        run = meta.add_run(f"Site Analisado: {data['url']}")
        run.font.italic = True
        run.font.size = Pt(10)

    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Processar o conteúdo da análise
    content = data.get("analysis", "")
    
    # Dividir o conteúdo por seções baseadas em numeração ou tópicos comuns
    lines = content.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        # Identificar Títulos (ex: 1. OBJETIVO, ou **TÍTULO**)
        if line.startswith(('1.', '2.', '3.', '4.', '5.')) or line.isupper() or (line.startswith('**') and line.endswith('**')):
            clean_title = line.replace('**', '')
            h = doc.add_heading(clean_title, level=1)
            h.runs[0].font.color.rgb = RGBColor(0, 51, 102) # Azul escuro profissional
        elif line.startswith('-') or line.startswith('*'):
            p = doc.add_paragraph(line, style='List Bullet')
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_after = Pt(6)

    # Rodapé
    doc.add_paragraph("\n\n")
    footer = doc.add_paragraph("Relatório gerado automaticamente pela Laura AI - Sua Assistente Estratégica.")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(9)
    footer.runs[0].font.italic = True

    # Salvar
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    full_path = os.path.join(BASE_DIR, filename)
    doc.save(full_path)
    return full_path

def get_organized_path(category="Geral"):
    """Cria e retorna o caminho organizado na pasta Downloads."""
    import os
    from datetime import datetime
    
    downloads_path = os.path.join(os.path.expanduser("~"), "Downloads")
    base_laura_path = os.path.join(downloads_path, "Laura_Relatorios")
    
    date_str = datetime.now().strftime("%Y-%m-%d")
    final_dir = os.path.join(base_laura_path, date_str, category)
    
    if not os.path.exists(final_dir):
        os.makedirs(final_dir, exist_ok=True)
        
    return final_dir

def execute(query, say, takeCommand, context=None):
    BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    client = context.get("client") if context else None
    model = context.get("model_to_use") if context else None

    # Detecta se é pedido de forecast/previsão
    forecast_kws = ["previsão", "forecast", "projetar", "prever", "pipeline", "funil", "receita"]
    is_forecast = any(k in query.lower() for k in forecast_kws)

    if is_forecast:
        say("Modo Forecast ativado. Me passe os dados do seu pipeline: quantos leads, ticket médio e em qual fase a maioria está?")
        pipeline_info = takeCommand(timeout=30, phrase_time_limit=60)
        if not pipeline_info or pipeline_info == "none":
            say("Tudo bem. Para gerar um forecast preciso, precisarei dos dados do seu funil de vendas.")
            return True

        say("Calculando sua previsão de receita e saúde do funil...")
        if client and model:
            hoje = datetime.datetime.now().strftime("%d/%m/%Y")
            prompt = (
                f"Você é Laura, especialista em Revenue Operations e Forecast de Vendas.\n"
                f"Data de hoje: {hoje}\n"
                f"Dados do pipeline informados: '{pipeline_info}'\n\n"
                f"Entregue uma análise completa de forecast:\n"
                f"1. SAÚDE DO PIPELINE: Avalie os dados e sinalize gargalos ou riscos.\n"
                f"2. TAXA DE CONVERSÃO ESTIMADA: Por fase (prospecção → qualificação → proposta → fechamento).\n"
                f"3. PREVISÃO DE RECEITA: Valor projetado para os próximos 30, 60 e 90 dias.\n"
                f"4. RISCO DE PERDA: Quais deals estão em risco e por quê.\n"
                f"5. AÇÕES PRIORITÁRIAS: Top 3 ações para aumentar a taxa de fechamento agora.\n\n"
                f"Use dados reais estimados. Seja objetivo. Comece com: 'Senhor, aqui está seu forecast de vendas:'"
            )
            try:
                response = client.chat.completions.create(
                    model=model,
                    messages=[{"role": "user", "content": prompt}]
                )
                resultado = response.choices[0].message.content
                say(resultado)
                # Salva o forecast para posterior export em Word
                forecast_file = os.path.join(BASE_DIR, "last_marketing_analysis.json")
                with open(forecast_file, "w", encoding="utf-8") as f:
                    json.dump({"url": "Pipeline Interno", "analysis": resultado}, f, ensure_ascii=False)
                say("Deseja que eu exporte este forecast em um relatório Word?")
                ans = takeCommand(timeout=10)
                if ans and any(w in ans.lower() for w in ["sim", "pode", "export", "word", "relatório"]):
                    from skills.report_generator import create_professional_report, get_organized_path
                    ts = datetime.datetime.now().strftime("%H%M")
                    fname = f"Forecast_Vendas_{ts}.docx"
                    tdir = get_organized_path("Forecast")
                    fpath = os.path.join(tdir, fname)
                    create_professional_report({"url": "Pipeline Interno", "analysis": resultado}, fpath)
                    say(f"Relatório de forecast salvo em Downloads, Laura Relatórios, Forecast.")
                    os.startfile(tdir)
            except Exception as e:
                say(f"Erro ao gerar forecast: {e}")
        return True

    analysis_file = os.path.join(BASE_DIR, "last_marketing_analysis.json")

    
    # Se estivermos em modo projeto, os dados podem vir direto no contexto
    data = context.get("report_data") if context else None
    
    # Se não houver dados, tenta carregar a última análise de marketing
    if not data and os.path.exists(analysis_file):
        try:
            with open(analysis_file, "r", encoding="utf-8") as f:
                data = json.load(f)
        except: pass

    # NOVO: Se ainda não houver dados, tenta usar o Histórico de Chat como fonte!
    if not data:
        history_file = os.path.join(BASE_DIR, "chat_history.json")
        if os.path.exists(history_file):
            try:
                print("[Report] Usando histórico de chat como fonte...")
                with open(history_file, "r", encoding="utf-8") as f:
                    history = json.load(f)
                # Pega as últimas 3 interações (User + Assistant)
                relevant_history = history[-4:] 
                content = "\n".join([f"{m['role'].upper()}: {m['content']}" for m in relevant_history])
                data = {
                    "url": "Conversa Recente",
                    "analysis": f"Relatório gerado a partir da conversa:\n\n{content}"
                }
            except: pass

    if not data:
        say("Senhor, não encontrei dados para gerar o relatório. Por favor, realize uma análise primeiro.")
        return True

    say("Processando e organizando seu relatório profissional...")
    
    try:
        timestamp = datetime.datetime.now().strftime("%H%M")
        site_name = data.get("url", "Projeto").split("//")[-1].split(".")[0]
        category = "Marketing" if "marketing" in query.lower() or "ads" in query.lower() else "Geral"
        
        filename = f"Relatorio_{site_name}_{timestamp}.docx"
        target_dir = get_organized_path(category)
        full_path = os.path.join(target_dir, filename)
        
        # Gerar o relatório
        create_professional_report(data, full_path)
        
        print(f"[Report] Salvo em: {full_path}")
        say(f"Relatório concluído com sucesso e organizado em: Downloads, Laura Relatórios, {category}.")
        
        # Se não for uma chamada automática do agente, perguntar se quer abrir
        is_auto = context.get("is_autonomous", False) if context else False
        if not is_auto:
            say("Deseja que eu abra a pasta para você agora?")
            ans = takeCommand(timeout=8)
            if ans and any(w in ans.lower() for w in ["sim", "pode", "abre", "ok", "abrir"]):
                os.startfile(target_dir)
                say("Abrindo pasta de destino.")
        else:
            # Em modo autônomo, apenas sinaliza conclusão
            return {"status": "success", "file": full_path}
            
    except Exception as e:
        print(f"Erro ao gerar/organizar DOCX: {e}")
        say("Houve um problema ao organizar o arquivo na pasta Downloads.")
    
    return True
