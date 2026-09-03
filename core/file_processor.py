"""
Laura AI - File Processor Module
Extrai texto de arquivos PDF, DOCX, XLSX e CSV para análise pela IA.
"""
import os
import csv
import io
import traceback


def extract_text(file_path: str, max_chars: int = 4000) -> dict:
    """
    Extrai o conteúdo textual de um arquivo suportado.
    
    Args:
        file_path: Caminho absoluto do arquivo.
        max_chars: Limite de caracteres extraídos (para caber no contexto da IA).
    
    Returns:
        dict com 'success', 'text', 'filename', 'type', e opcionalmente 'error'.
    """
    if not os.path.exists(file_path):
        return {"success": False, "error": "Arquivo não encontrado.", "text": ""}

    filename = os.path.basename(file_path)
    ext = os.path.splitext(filename)[1].lower()

    extractors = {
        ".pdf": _extract_pdf,
        ".docx": _extract_docx,
        ".doc": _extract_docx,  # python-docx suporta .doc parcialmente
        ".xlsx": _extract_xlsx,
        ".xls": _extract_xlsx,
        ".csv": _extract_csv,
    }

    extractor = extractors.get(ext)
    if not extractor:
        return {
            "success": False,
            "error": f"Formato '{ext}' não suportado. Use PDF, DOCX, XLSX ou CSV.",
            "text": "",
        }

    try:
        text = extractor(file_path)
        # Truncar se necessário
        truncated = False
        if len(text) > max_chars:
            text = text[:max_chars]
            truncated = True

        return {
            "success": True,
            "text": text,
            "filename": filename,
            "type": ext.replace(".", "").upper(),
            "truncated": truncated,
            "original_length": len(text) if not truncated else ">" + str(max_chars),
        }
    except ImportError as e:
        lib_name = str(e).split("'")[-2] if "'" in str(e) else str(e)
        return {
            "success": False,
            "error": f"Biblioteca necessária não instalada: {lib_name}. Execute: pip install {lib_name}",
            "text": "",
        }
    except Exception as e:
        traceback.print_exc()
        return {
            "success": False,
            "error": f"Erro ao processar arquivo: {str(e)}",
            "text": "",
        }


def _extract_pdf(file_path: str) -> str:
    """Extrai texto de arquivo PDF usando PyPDF2."""
    from PyPDF2 import PdfReader

    reader = PdfReader(file_path)
    pages = []
    for i, page in enumerate(reader.pages):
        page_text = page.extract_text()
        if page_text:
            pages.append(f"--- Página {i + 1} ---\n{page_text.strip()}")
    
    if not pages:
        return "[PDF sem texto extraível — pode conter apenas imagens]"
    return "\n\n".join(pages)


def _extract_docx(file_path: str) -> str:
    """Extrai texto de arquivo DOCX usando python-docx."""
    from docx import Document

    doc = Document(file_path)
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Também extrair conteúdo de tabelas
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.replace("|", "").strip():
                paragraphs.append(row_text)

    if not paragraphs:
        return "[Documento vazio ou sem texto extraível]"
    return "\n".join(paragraphs)


def _extract_xlsx(file_path: str) -> str:
    """Extrai texto de arquivo XLSX usando openpyxl."""
    from openpyxl import load_workbook

    wb = load_workbook(file_path, read_only=True, data_only=True)
    sheets_text = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows = []
        for row in ws.iter_rows(values_only=True):
            row_values = [str(cell) if cell is not None else "" for cell in row]
            # Pular linhas totalmente vazias
            if any(v.strip() for v in row_values):
                rows.append(" | ".join(row_values))
        
        if rows:
            sheets_text.append(f"--- Planilha: {sheet_name} ---\n" + "\n".join(rows))

    wb.close()

    if not sheets_text:
        return "[Planilha vazia]"
    return "\n\n".join(sheets_text)


def _extract_csv(file_path: str) -> str:
    """Extrai texto de arquivo CSV."""
    rows = []
    # Tentar detectar encoding
    encodings = ["utf-8", "latin-1", "cp1252"]
    
    for encoding in encodings:
        try:
            with open(file_path, "r", encoding=encoding, newline="") as f:
                reader = csv.reader(f)
                for row in reader:
                    row_text = " | ".join(cell.strip() for cell in row)
                    if row_text.replace("|", "").strip():
                        rows.append(row_text)
            break  # Encoding funcionou
        except UnicodeDecodeError:
            continue

    if not rows:
        return "[CSV vazio ou com encoding não suportado]"
    return "\n".join(rows)


def get_supported_extensions() -> list:
    """Retorna lista de extensões suportadas para o diálogo de arquivo."""
    return [
        "Documentos suportados (*.pdf;*.docx;*.doc;*.xlsx;*.xls;*.csv)",
        "PDF (*.pdf)",
        "Word (*.docx;*.doc)",
        "Excel (*.xlsx;*.xls)",
        "CSV (*.csv)",
    ]
